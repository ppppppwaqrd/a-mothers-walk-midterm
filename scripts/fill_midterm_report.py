"""Fill the official midterm Word template (Copy of รายงานโครงงานเกม 2D.docx)."""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

GAME = Path(r"c:\Users\jakkr\OneDrive\เดสก์ท็อป\Gamedev\lab4\2D-Platformer-Starter-Kit-main")
FOLDER = Path(r"c:\Users\jakkr\OneDrive\เดสก์ท็อป\Gamedev\lab4\MIdterm project")
SRC = FOLDER / "Copy of รายงานโครงงานเกม 2D.docx"
OUT_NAMED = FOLDER / "รายงานโครงงานเกม_กล่องข้าวน้อย.docx"
SHOTS = GAME / "scripts" / "_shots"
DOCS = GAME / "docs"


def _sarabun(run, size_pt: float = 16, bold: bool | None = None) -> None:
    run.font.name = "Sarabun"
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), "Sarabun")


def set_text(para: Paragraph, text: str, bold: bool | None = None) -> None:
    if not para.runs:
        run = para.add_run(text)
        _sarabun(run, bold=bold)
        return
    para.runs[0].text = text
    if bold is not None:
        para.runs[0].bold = bold
    for extra in para.runs[1:]:
        extra.text = ""


def insert_after(paragraph: Paragraph, text: str, bold: bool | None = None) -> Paragraph:
    new_elm = deepcopy(paragraph._p)
    paragraph._p.addnext(new_elm)
    new_para = Paragraph(new_elm, paragraph._parent)
    set_text(new_para, text, bold=bold)
    return new_para


def insert_many(after: Paragraph, texts: list[str]) -> Paragraph:
    cur = after
    for text in texts:
        cur = insert_after(cur, text)
    return cur


def insert_picture(after: Paragraph, path: Path, caption: str, width: float = 5.5) -> Paragraph:
    cur = after
    if caption:
        cur = insert_after(cur, caption)
    pic = insert_after(cur, "")
    run = pic.add_run()
    run.add_picture(str(path), width=Inches(width))
    return pic


def find_startswith(doc: Document, prefix: str) -> Paragraph:
    for para in doc.paragraphs:
        if para.text.strip().startswith(prefix):
            return para
    raise KeyError(prefix)


def next_paragraph(paragraph: Paragraph) -> Paragraph:
    nxt = paragraph._p.getnext()
    if nxt is None:
        raise RuntimeError("no next paragraph")
    return Paragraph(nxt, paragraph._parent)


def main() -> None:
    doc = Document(str(SRC))

    set_text(doc.paragraphs[2], "วิชา [  ] CP352203  Computer Game Development")
    set_text(doc.paragraphs[3], "วิชา [x] CP410844  Computer Game Development")

    name_para = find_startswith(doc, "ชื่อเกม")
    set_text(name_para, "ชื่อเกม :  A Mother's Walk (กล่องข้าวน้อย)", bold=False)
    if name_para.runs:
        name_para.runs[0].bold = False
        # Keep the label bold like the template.
        name_para.runs[0].text = ""
        label = name_para.add_run("ชื่อเกม :  ")
        _sarabun(label, bold=True)
        value = name_para.add_run("A Mother's Walk (กล่องข้าวน้อย)")
        _sarabun(value, bold=False)

    group_para = find_startswith(doc, "กลุ่มที่")
    set_text(group_para, "กลุ่มที่  4", bold=True)

    members = [
        "673380348-4    นายสรวิศ สำราญบึงแก    สาขาวิชาวิทยาการคอมพิวเตอร์ (ปัญญาประดิษฐ์)",
        "673380308-6    นายจักรภัทร เวียงสิมมา    สาขาวิชาวิทยาการคอมพิวเตอร์ (ปัญญาประดิษฐ์)",
        "673380328-0    นายพรหมพัฒน ศิริภัคกุลวัฒน์    สาขาวิชาวิทยาการคอมพิวเตอร์ (ปัญญาประดิษฐ์)",
    ]
    m1 = find_startswith(doc, "<รหัสนักศึกษา")
    set_text(m1, members[0])
    m2 = next_paragraph(m1)
    set_text(m2, members[1])
    m3 = insert_after(m2, members[2])

    # --- Theme ---
    theme_h = find_startswith(doc, "ธีม")
    theme_p = next_paragraph(theme_h)
    set_text(
        theme_p,
        "เกมแนว 2D Platformer / Adventure ที่ตีความนิทานพื้นบ้านไทย “กล่องข้าวน้อยฆ่าแม่” "
        "ใหม่ในมุมของมารดา ผู้เล่นไม่ได้เห็นตอนจบโศกนาฏกรรมของนิทาน แต่ต้องพาแม่เดินทาง "
        "นำกล่องข้าวไปส่งลูกชายให้ทันก่อนที่ความอดทนของลูกจะหมด",
    )
    insert_many(
        theme_p,
        [
            "โทนศิลปะเป็นพิกเซลอาร์ตชนบทอีสาน นำเสนอผ่านหน้าจอแบบสมุดนิทาน "
            "(พลิกหน้ากระดาษ บทนำเรื่อง และ HUD หมึกบนกระดาษ) พื้นหลังใช้ภาพสีน้ำแบบพารัลแลกซ์แยกตามเวลาของวันในแต่ละด่าน",
            "พัฒนาด้วย Godot Engine 4.7 ต่อยอดจาก 2D Platformer Starter Kit ของรายวิชา "
            "แล้วออกแบบระบบ ศิลปะ เสียง ด่าน และกติกาใหม่ทั้งเกม",
        ],
    )

    # --- Audience ---
    aud_h = find_startswith(doc, "กลุ่มเป้าหมาย")
    aud_p = next_paragraph(aud_h)
    set_text(
        aud_p,
        "นักเรียน นักศึกษา และผู้เล่นทั่วไปที่สนใจเกมสั้นเล่นง่าย อายุประมาณ 12 ปีขึ้นไป "
        "รวมถึงผู้ที่อยากเล่นนิทานพื้นบ้านในรูปแบบอินเทอร์แอคทีฟ รองรับทั้งคีย์บอร์ดและปุ่มบนจอสัมผัส",
    )

    # --- Story ---
    story_h = find_startswith(doc, "เนื้อเรื่องย่อ")
    story_p = next_paragraph(story_h)
    set_text(
        story_p,
        "ในหมู่บ้านอีสาน แม่หุงข้าวเหนียวใส่กล่องใบน้อย เพื่อนำไปให้ลูกชายชื่อไอ้ทองที่รออยู่ปลายทาง "
        "ท้องของลูกว่างเปล่า และเวลาไม่คอยใคร แม่จึงออกเดินจากหมู่บ้าน ผ่านป่าไผ่ ทางขรุขระ ทุ่งนา "
        "และคูน้ำกลางคืน ระหว่างทางมีหนาม หอกไผ่ สัตว์ป่า ชาวบ้านที่ขอความช่วยเหลือ "
        "ศาลเทวดาที่ลองใจ และปริศนาหินลูกรังที่ต้องดันทับสวิตช์จึงจะข้ามไปได้",
    )
    insert_after(
        story_p,
        "เป้าหมายของผู้เล่นคือส่งกล่องข้าวให้ถึงมืออ้ายทองก่อนหลอดความอดทนหมด "
        "และเก็บกระติบข้าวให้ครบ 8 ใบตลอดเรื่อง หากไปไม่ทัน เรื่องจบที่หน้า “อ้ายทองรอไม่ไหว” "
        "หากถึงแต่กระติบไม่ครบ แม่ถึงข้างลูกแล้วแต่ข้าวยังไม่พอให้ยิ้มได้ "
        "ถ้าเก็บครบและส่งถึง เรื่องจบแบบสุขที่กล่องข้าวยังอุ่น",
    )

    # --- Rules ---
    rules_h = find_startswith(doc, "รูปแบบการเล่น")
    rules_p = next_paragraph(rules_h)
    set_text(
        rules_p,
        "เริ่มเกมใหม่จะเปิดสมุดนิทาน 6 หน้า (เนื้อเรื่องและกติกาทั้งเกม) ก่อนเข้าด่าน 1 "
        "เลือกเล่นต่อหรือลองด่านอีกครั้งจะข้ามบทนำนี้ เกมมี 6 ด่านต่อเนื่อง "
        "แต่ละด่านเปิดด้วยหน้าบทพร้อมคำใบ้ “เก็บกระติบให้ครบ 8 ใบ · อย่าให้หลอดไอ้ทองหมด”",
    )
    insert_many(
        rules_p,
        [
            "การควบคุม: A หรือ ลูกศรซ้าย เดินซ้าย · D หรือ ลูกศรขวา เดินขวา · Space หรือ W กระโดด · "
            "J ปาหิน (ก้อนหินมีจำกัด ต้องเก็บเติม) · Esc หรือ P เปิดหน้าพัก ปรับเสียงได้ · "
            "บนจอสัมผัสใช้ปุ่มมุมล่างแทนได้",
            "หัวใจบน HUD คือจำนวนชีวิต เสียชีวิตแล้วเกิดใหม่ที่ศาลเซฟจุดล่าสุดในด่านนั้น "
            "หัวใจหมดทุกดวงถือว่าแพ้ แถบเขียวคือเลือดของแม่ แถบน้ำตาลคือความอดทนของไอ้ทอง "
            "ตัวเลขกระติบคือข้าวที่เก็บได้ และตัวเลขหินคือกระสุน",
            "หลอดความอดทนของไอ้ทองลดลงตลอดเวลาขณะเล่นด่าน เมื่อหมดจะแพ้ทันที "
            "แม้เลือดของแม่ยังเหลือ ช่วยชาวบ้านหรือเล่นมินิเกมศาลเทวดาผ่าน จะได้ความอดทนและก้อนหินเพิ่ม "
            "ถ้ามินิเกมไม่ผ่าน หลอดไอ้ทองจะลดเร็วขึ้นในด่านนั้น",
            "ศัตรู: งูและนกกาปาหินไม่กี่ครั้งก็พอ หมูป่าต้องปา 3 ครั้ง ชนแล้วเลือดเหลือประมาณครึ่ง "
            "ควายต้องปา 5 ครั้ง ชนแล้วตายทันที กับดักหนาม หอกไผ่ ลูกตุ้ม ใบมีด และตกคูน้ำทำให้เสียเลือดหรือตาย",
            "บางจุดต้องดันหินลูกรังไปทับสวิตช์ เพื่อเปิดกำแพงไผ่หรือวางสะพาน จึงจะเดินต่อได้ "
            "ประตูจบด่านยืนบนพื้นถึงจะเข้าได้ เมื่อถึงด่านสุดท้ายและส่งกล่องข้าวสำเร็จ "
            "เกมจะนับกระติบ 8 ใบเพื่อเลือกตอนจบสุขหรือตอนจบที่ข้าวไม่พอ",
        ],
    )

    # --- Characters ---
    char_h = find_startswith(doc, "ตัวละคร")
    char_p = next_paragraph(char_h)
    set_text(char_p, "แม่ (ผู้เล่น) — นางเอก ถือกระติ๊บข้าว เดิน กระโดด ปาหิน และดันของ เพื่อนำอาหารไปส่งลูก")
    insert_many(
        char_p,
        [
            "ไอ้ทอง — ลูกชายที่รออยู่กลางทุ่ง ไม่ได้อยู่บนจอระหว่างเดินทาง แต่หลอดความอดทนของเขาคือตัวจับเวลาทั้งเกม",
            "ชาวบ้าน — NPC ตามทาง เดินเข้าไปช่วยได้หนึ่งครั้ง จะได้ความอดทนและก้อนหิน",
            "เทวดาประจำศาล — เปิดมินิเกมลองใจแม่ (ไล่อีกาจากนา / ไล่ควายออกจากลาน / อย่าให้ไก่คุ้ยข้าวตาก)",
            "งู หมูป่า นกกา ควาย — ศัตรูตามเส้นทาง ความทนและโทษตอนชนต่างกันตามชนิด",
        ],
    )

    # --- Items ---
    item_h = find_startswith(doc, "ไอเทม")
    item_p = next_paragraph(item_h)
    set_text(item_p, "กระติบข้าว — ไอเทมหลัก นับเข้าโควตา 8 ใบสำหรับตอนจบสุข (ของเก็บอย่างอื่นไม่นับโควตานี้)")
    insert_many(
        item_p,
        [
            "ก้อนหิน — เติมกระสุนปา เริ่มด่านมีจำกัด สูงสุด 10 ก้อน",
            "หัวใจ — ฟื้นเลือดของแม่ บางดวงเพิ่มชีวิต",
            "น้ำเต้าความเร็ว / ใบพลู — บูสต์ความเร็วหรือการกระโดดชั่วคราว",
            "จุดพักศาล (สีเขียว) — บันทึกจุดเกิดในด่านนั้น ตายแล้วยังอยู่ในด่านเดิม",
            "ศาลเทวดา — มินิเกม เล่นผ่านได้ความอดทน +20 และก้อนหิน +3 ไม่ผ่านแล้วหลอดไอ้ทองลดเร็วขึ้น",
            "หินลูกรัง + สวิตช์ + กำแพงไผ่ / สะพานไม้ — ปริศนาดันของ มีในด่าน 1, 4 และ 5",
            "แพเลื่อน ลิฟต์ กระดานดีด — ช่วยข้ามช่องว่างหรือขึ้นที่สูง",
            "กับดัก: หนาม หอกไผ่ ลูกตุ้มหิน ใบมีด และคูน้ำ (ตกแล้วตาย)",
        ],
    )

    # --- Levels + screenshots ---
    lv_h = find_startswith(doc, "ด่าน/ฉาก")
    lv_p = next_paragraph(lv_h)
    set_text(
        lv_p,
        "เกมมี 6 ด่าน เรียงตามเวลาของวันจากเช้าถึงแสงสุดท้าย แต่ละด่านมีพื้นหลังพารัลแลกซ์และเพลงของตนเอง "
        "ภาพด้านล่างเป็นหน้าจอจริงจากเกม",
    )
    cur = lv_p

    shots: list[tuple[Path, str, str]] = [
        (SHOTS / "menu.png", "ภาพที่ 1  หน้าเมนูหลัก ธีมสมุดนิทาน", ""),
        (
            SHOTS / "level_01.png",
            "ภาพที่ 2  ด่าน 1 ออกจากหมู่บ้าน (เช้าตรู่)",
            "สอนการเดิน กระโดด ปาหิน เก็บกระติบ และดันหินลูกรังทับสวิตช์เปิดกำแพงไผ่ มีงู หนาม ชาวบ้านที่ขอความช่วยเหลือ และจุดพักศาล",
        ),
        (
            SHOTS / "level_02.png",
            "ภาพที่ 3  ด่าน 2 ป่าไผ่ (สาย)",
            "มีแพเลื่อน อีกา และศาลเทวดามินิเกมไล่อีกาจากนาข้าว",
        ),
        (
            SHOTS / "level_03.png",
            "ภาพที่ 4  ด่าน 3 ทางขรุขระ (เที่ยง)",
            "ทางสูงมีลิฟต์ ควายที่ชนแล้วตายทันที ลูกตุ้มหิน และมินิเกมไล่ควายออกจากลาน",
        ),
        (
            SHOTS / "level_04.png",
            "ภาพที่ 5  ด่าน 4 ทุ่งนากลางทาง (บ่ายแก่)",
            "ทุ่งนามีกระดานดีด ปริศนาสะพาน และมินิเกมอย่าให้ไก่คุ้ยข้าวตาก",
        ),
        (
            SHOTS / "level_05.png",
            "ภาพที่ 6  ด่าน 5 คูน้ำกลางคืน (พลบค่ำ)",
            "อันตรายหนาแน่นขึ้น มีหอกไผ่ ใบมีด ลูกตุ้ม อีกา และต้องดันหินเปิดสะพานข้ามคู ตกน้ำแล้วตาย",
        ),
        (
            SHOTS / "level_06.png",
            "ภาพที่ 7  ด่าน 6 ส่งกล่องข้าวให้อ้ายทอง (แสงสุดท้าย)",
            "ด่านปิดเรื่อง พาแม่ถึงลูกแล้วเกมนับกระติบเพื่อเลือกตอนจบ",
        ),
        (SHOTS / "game_win.png", "ภาพที่ 8  หน้าชนะ (ตอนจบสุขเมื่อกระติบครบ 8 ใบ)", ""),
        (SHOTS / "game_over.png", "ภาพที่ 9  หน้าแพ้ เมื่อชีวิตหมด หรือเมื่อหลอดไอ้ทองหมด", ""),
    ]
    for path, caption, extra in shots:
        if not path.exists():
            alt = DOCS / path.name
            path = alt if alt.exists() else path
        if path.exists():
            cur = insert_picture(cur, path, caption)
        else:
            cur = insert_after(cur, caption + " (ยังไม่มีไฟล์ภาพ)")
        if extra:
            cur = insert_after(cur, extra)

    demo1, demo2 = DOCS / "demo1.jpg", DOCS / "demo2.jpg"
    if demo1.exists() and demo2.exists():
        cur = insert_after(cur, "ภาพเพิ่มเติมจากหน้าเมนูและหน้าเล่นที่ใช้ประกอบการส่งงาน")
        cur = insert_picture(cur, demo1, "")
        cur = insert_picture(cur, demo2, "")

    # --- Benefits + links (assignment also asks GitHub + itch.io) ---
    ben_h = find_startswith(doc, "ประโยชน์ของเกม")
    ben_p = next_paragraph(ben_h)
    set_text(
        ben_p,
        "เกมนำนิทานพื้นบ้านมาเล่าใหม่แบบอินเทอร์แอคทีฟ ให้ผู้เล่นสวมบทมารดาและเห็นผลของการรีบหรือการแวะช่วยคน "
        "ผ่านหลอดความอดทน โควตากระติบ และตอนจบสองแบบ จึงฝึกทั้งการบริหารเวลาและการแก้ปริศนาสั้น ๆ",
    )
    last = insert_many(
        ben_p,
        [
            "ผู้เล่นได้ฝึกการควบคุมเกมแพลตฟอร์ม การสังเกตกับดัก และการตัดสินใจว่าจะเร่งไปส่งข้าวหรือแวะศาลเทวดา",
            "เป็นผลงานการเรียนรู้การพัฒนาเกม 2D ด้วย Godot ของนักศึกษา "
            "(ฉาก ระบบ ศิลปะ เสียง การส่งออกเว็บ และการจัดทำเอกสารโครงงาน)",
        ],
    )
    link_h = insert_after(last, "ลิงก์ผลงาน", bold=True)
    insert_many(
        link_h,
        [
            "GitHub Project (ซอร์สโค้ด):  https://github.com/ppppppwaqrd/a-mothers-walk-midterm",
            "GitHub Pages (เล่นบนเว็บ):  https://ppppppwaqrd.github.io/a-mothers-walk-midterm/",
            "itch.io (เผยแพร่เกม HTML5 ตามข้อกำหนดรายวิชา):  อัปโหลดจากโฟลเดอร์ itch/ ของโปรเจกต์ แล้ววาง URL เกมที่นี่  ______________________________",
            "เครื่องมือ: Godot Engine 4.7  ·  ฟอนต์ Charmonman / Mali (SIL OFL)  ·  เพลงและเสียงเอฟเฟกต์สร้างสำหรับโปรเจกต์นี้",
        ],
    )

    for dest in (SRC, OUT_NAMED):
        doc.save(str(dest))

    log = GAME / "_report_filled.txt"
    lines = [f"{i:03d}|{p.text}" for i, p in enumerate(doc.paragraphs)]
    log.write_text("\n".join(lines), encoding="utf-8")
    print("saved filled report, paragraphs:", len(doc.paragraphs))


if __name__ == "__main__":
    main()
